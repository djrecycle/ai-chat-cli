"""Load local documents into chat context."""

from __future__ import annotations

import base64
import mimetypes
from dataclasses import dataclass
from pathlib import Path

from .providers import ChatMessage, ImageAttachment


MAX_DOCUMENT_CHARS = 60_000
MAX_INLINE_IMAGE_BYTES = 15 * 1024 * 1024
IMAGE_EXTENSIONS = {
    ".bmp",
    ".gif",
    ".jpeg",
    ".jpg",
    ".png",
    ".tif",
    ".tiff",
    ".webp",
}


@dataclass(frozen=True)
class LoadedDocument:
    path: Path
    content: str
    truncated: bool = False


class DocumentLoadError(Exception):
    """Raised when a document cannot be loaded as text."""


def load_document(path_text: str, *, max_chars: int = MAX_DOCUMENT_CHARS) -> LoadedDocument:
    path = Path(path_text).expanduser()
    if not path.exists():
        raise DocumentLoadError(f"File tidak ditemukan: {path}")
    if not path.is_file():
        raise DocumentLoadError(f"Path bukan file: {path}")

    suffix = path.suffix.lower()
    if suffix == ".pdf":
        content = _read_pdf(path)
    elif suffix == ".docx":
        content = _read_docx(path)
    elif suffix in IMAGE_EXTENSIONS:
        content = _read_image(path)
    else:
        content = _read_text(path, max_chars=max_chars + 1)

    content = content.strip()
    if not content:
        raise DocumentLoadError(f"File kosong atau tidak berisi teks yang bisa dibaca: {path}")

    truncated = len(content) > max_chars
    if truncated:
        content = content[:max_chars].rstrip()

    return LoadedDocument(path=path, content=content, truncated=truncated)


def build_document_prompt(document: LoadedDocument, question: str | None = None) -> str:
    suffix = "\n\n[Catatan: isi file dipotong karena terlalu panjang.]" if document.truncated else ""
    instruction = question.strip() if question and question.strip() else "Tolong ringkas isi dokumen ini."
    return (
        f"Saya melampirkan isi file `{document.path}`.\n\n"
        f"```text\n{document.content}\n```"
        f"{suffix}\n\n"
        f"Pertanyaan/instruksi saya: {instruction}"
    )


def build_image_message(path_text: str, question: str | None = None) -> ChatMessage:
    """Build a multimodal message containing the original image bytes."""
    path = Path(path_text).expanduser()
    if not path.exists():
        raise DocumentLoadError(f"File tidak ditemukan: {path}")
    if not path.is_file():
        raise DocumentLoadError(f"Path bukan file: {path}")
    if path.suffix.lower() not in IMAGE_EXTENSIONS:
        raise DocumentLoadError(f"File bukan format gambar yang didukung: {path}")

    try:
        raw = path.read_bytes()
    except OSError as exc:
        raise DocumentLoadError(f"Gagal membaca gambar: {path}") from exc
    if not raw:
        raise DocumentLoadError(f"File gambar kosong: {path}")
    if len(raw) > MAX_INLINE_IMAGE_BYTES:
        size_mb = len(raw) / (1024 * 1024)
        raise DocumentLoadError(
            f"Gambar terlalu besar ({size_mb:.1f} MB). Maksimum 15 MB untuk upload langsung."
        )

    try:
        from PIL import Image

        with Image.open(path) as image:
            image.verify()
    except ImportError as exc:
        raise DocumentLoadError("Validasi gambar membutuhkan paket `Pillow`.") from exc
    except OSError as exc:
        raise DocumentLoadError(f"File bukan gambar valid: {path}") from exc

    mime_type = mimetypes.guess_type(path.name)[0] or "application/octet-stream"
    prompt = question.strip() if question and question.strip() else (
        "Jelaskan secara menyeluruh apa yang terlihat pada gambar ini."
    )
    visible_content = f"[Gambar: {path.name}]\n{prompt}"
    attachment = ImageAttachment(
        name=path.name,
        mime_type=mime_type,
        data=base64.b64encode(raw).decode("ascii"),
    )
    return ChatMessage("user", visible_content, images=[attachment])


def _read_text(path: Path, *, max_chars: int) -> str:
    raw = path.read_bytes()
    if b"\x00" in raw[:4096]:
        raise DocumentLoadError(
            f"File terlihat seperti binary dan tidak bisa dibaca sebagai teks: {path}"
        )

    data = raw[: max_chars * 4]
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentLoadError(f"Gagal decode file teks: {path}")


def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        try:
            from PyPDF2 import PdfReader
        except ImportError as exc:
            raise DocumentLoadError(
                "Membaca PDF butuh paket tambahan. Install salah satu: "
                "`pip install pypdf` atau `pip install PyPDF2`."
            ) from exc

    try:
        reader = PdfReader(str(path))
        if reader.is_encrypted:
            try:
                unlocked = reader.decrypt("")
            except Exception as exc:
                raise DocumentLoadError(
                    f"PDF dilindungi kata sandi dan tidak bisa dibaca: {path}"
                ) from exc
            if not unlocked:
                raise DocumentLoadError(
                    f"PDF dilindungi kata sandi dan tidak bisa dibaca: {path}"
                )

        pages = [page.extract_text() or "" for page in reader.pages]
    except DocumentLoadError:
        raise
    except Exception as exc:
        raise DocumentLoadError(f"Gagal membaca PDF: {path}. {exc}") from exc
    text = "\n\n".join(pages).strip()
    if not text:
        raise DocumentLoadError(
            "PDF tidak berisi teks yang bisa diekstrak. Jika ini PDF hasil scan, "
            "ubah halaman menjadi gambar lalu pakai `/file gambar.png ...`, atau "
            "jalankan OCR PDF di luar aplikasi terlebih dahulu."
        )
    return text


def _read_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise DocumentLoadError(
            "Membaca DOCX butuh paket tambahan: `pip install python-docx`."
        ) from exc

    try:
        doc = Document(str(path))
    except Exception as exc:
        raise DocumentLoadError(f"Gagal membaca DOCX: {path}. {exc}") from exc

    parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _read_image(path: Path) -> str:
    try:
        from PIL import Image
    except ImportError as exc:
        raise DocumentLoadError(
            "Membaca gambar butuh paket tambahan: `pip install Pillow pytesseract`."
        ) from exc

    try:
        import pytesseract
    except ImportError as exc:
        raise DocumentLoadError(
            "Membaca teks dari gambar butuh OCR: `pip install pytesseract` "
            "dan install aplikasi Tesseract di sistem."
        ) from exc

    try:
        image = Image.open(path)
    except OSError as exc:
        raise DocumentLoadError(f"Gagal membuka gambar: {path}") from exc

    try:
        with image:
            try:
                text = pytesseract.image_to_string(image, lang="ind+eng")
            except pytesseract.TesseractError:
                text = pytesseract.image_to_string(image)
    except pytesseract.TesseractNotFoundError as exc:
        raise DocumentLoadError(
            "Aplikasi Tesseract belum ditemukan. Install dulu, misalnya: "
            "`sudo apt install tesseract-ocr tesseract-ocr-ind`."
        ) from exc
    except Exception as exc:
        raise DocumentLoadError(f"Gagal menjalankan OCR pada gambar: {path}. {exc}") from exc

    text = text.strip()
    if not text:
        raise DocumentLoadError(
            "Tidak ada teks yang terdeteksi di gambar. Pastikan gambar cukup jelas "
            "atau gunakan gambar dengan resolusi lebih tinggi."
        )
    return text
